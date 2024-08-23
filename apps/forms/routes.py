# -*- encoding: utf-8 -*-
"""
Copyright (c) 2019 - present AppSeed.us
"""

from flask import render_template, redirect, request, url_for, flash, session, send_file
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from apps import db, login_manager
from apps.forms import blueprint
from apps.forms.forms import CreateForm, TableForm
from apps.forms.models import Forms, TableEntry
from flask_login import current_user, login_required
from htmldocx import HtmlToDocx
from apps.decorators import subscription_required, proposal_charges_required, prevent_step_one_if_editing

@blueprint.route('/view_forms')
@login_required
@subscription_required
def view_forms():
    forms = Forms.query.filter_by(user_id=current_user.id).all()
    return render_template('forms/view_forms.html', forms=forms)

@blueprint.route('/new_form')
@login_required
@proposal_charges_required
def new_form():
    session.pop('form_data', None)
    return redirect(url_for('forms_blueprint.submit_form', step=1))

@blueprint.route('/edit_form/<int:form_id>/<int:step>', methods=['GET', 'POST'])
@login_required
@proposal_charges_required
def edit_form(form_id, step):
    form_in_question = Forms.query.get(form_id)
    if form_in_question.number_of_downloads == 0:
        return redirect(url_for('forms_blueprint.submit_form', step=1, form_id = form_id))
    else:
        return redirect(url_for('forms_blueprint.submit_form', step=2, form_id = form_id))

@blueprint.route('/submit_form/<int:step>', methods=['GET', 'POST'])
@login_required
@proposal_charges_required
@prevent_step_one_if_editing
def submit_form(step):
    form = CreateForm()
    table_form = TableForm()
    total_steps = 5

    if 'form_data' not in session:
        session['form_data'] = {}

    if step < 1 or step > total_steps:
        flash('Invalid step. Redirecting to the first step.', 'warning')
        return redirect(url_for('forms_blueprint.submit_form', step=1))

    form_id = request.args.get('form_id')
    if form_id:
        existing_form = Forms.query.get(form_id)
        if existing_form and existing_form.user_id == current_user.id:
            form_data = existing_form.__dict__.copy()
            form_data.pop('_sa_instance_state', None)
            session['form_data'] = form_data
            session.modified = True

    if request.method == 'POST':
        if step == 1:
            if form.validate_on_submit():
                form_data = {field.name: field.data for field in form if field.name not in ('csrf_token', 'submit', 'technical_approach_documentation', 'past_performance')}
                session['form_data'].update(form_data)
                session.modified = True
                form_id = session['form_data'].get('id')
                
                if form_id:
                    existing_form = Forms.query.get(form_id)
                    if existing_form and existing_form.user_id == current_user.id:
                        for key, value in form_data.items():
                            setattr(existing_form, key, value)
                    else:
                        flash('Error: Form not found or you do not have permission to edit it.', 'error')
                        return redirect(url_for('forms_blueprint.submit_form', step=1))
                else:
                    new_form = Forms(user_id=current_user.id, **form_data)
                    db.session.add(new_form)

                try:
                    db.session.commit()
                    if not form_id:
                        session['form_data']['id'] = new_form.id
                    session.modified = True
                    flash('Form data saved successfully.', 'success')
                except Exception as e:
                    db.session.rollback()
                    flash(f'Error saving form data: {str(e)}', 'error')
                    return redirect(url_for('forms_blueprint.submit_form', step=1, form_id = form_id))

                return redirect(url_for('forms_blueprint.submit_form', step=2, form_id = form_id))
            
        elif step == 2:
            if table_form.validate_on_submit():
                item_no = "{:03d}".format(table_form.item_no.data)
                new_entry = TableEntry(
                    form_id=session['form_data']['id'],
                    item_no=item_no,
                    description=table_form.description.data,
                    quantity=table_form.quantity.data,
                    unit=table_form.unit.data,
                    unit_price=table_form.unit_price.data
                )
                db.session.add(new_entry)
                db.session.commit()
                flash('Table entry added successfully.', 'success')
                return redirect(url_for('forms_blueprint.submit_form', step=2, form_id = form_id))
            else:
                flash('Error in form submission.', 'danger')
                return redirect(url_for('forms_blueprint.submit_form', step=1, form_id = form_id))

        elif step == 3:
            form_id = session['form_data'].get('id')
            if form_id:
                existing_form = Forms.query.get(form_id)
                if existing_form and existing_form.user_id == current_user.id:
                    existing_form.technical_approach_documentation = form.technical_approach_documentation.data
                    db.session.commit()
                    flash('Technical Approach Documentation saved successfully.', 'success')
                    return redirect(url_for('forms_blueprint.submit_form', step=4, form_id = form_id))
                else:
                    flash('Error: Form not found or you do not have permission to edit it.', 'error')
                    return redirect(url_for('forms_blueprint.submit_form', step=1, form_id = form_id))
            else:
                flash('Form ID not found in session.', 'error')
                return redirect(url_for('forms_blueprint.submit_form', step=1, form_id = form_id))

        elif step == 4:
            form_id = session['form_data'].get('id')
            if form_id:
                existing_form = Forms.query.get(form_id)
                if existing_form and existing_form.user_id == current_user.id:
                    existing_form.past_performance = form.past_performance.data
                    db.session.commit()
                    flash('Past Performance saved successfully.', 'success')
                    return redirect(url_for('forms_blueprint.submit_form', step=5, form_id = form_id))
                else:
                    flash('Error: Form not found or you do not have permission to edit it.', 'error')
                    return redirect(url_for('forms_blueprint.submit_form', step=1, form_id = form_id))
            else:
                flash('Form ID not found in session.', 'error')
                return redirect(url_for('forms_blueprint.submit_form', step=1, form_id = form_id))

        elif step == 5:
            form_id = session['form_data'].get('id')
            if form_id:
                flash('Form submitted successfully! Preparing your document.', 'success')
                return redirect(url_for('forms_blueprint.download_form', form_id = form_id))
            else:
                flash('Form ID not found in session.', 'error')
                return redirect(url_for('forms_blueprint.submit_form', step=1, form_id = form_id))

    if step == 1:
        for field in form:
            if field.name in session.get('form_data', {}):
                field.data = session['form_data'][field.name]

    table_entries = []
    if step == 2:
        form_id = session['form_data'].get('id')
        if form_id:
            table_entries = TableEntry.query.filter_by(form_id=form_id).all()
        for field in table_form:
            if field.name in session.get('form_data', {}):
                field.data = session['form_data'][field.name]

    if step == 3:
        form_id = session['form_data'].get('id')
        if form_id:
            existing_form = Forms.query.get(form_id)
            if existing_form and existing_form.user_id == current_user.id:
                form.technical_approach_documentation.data = existing_form.technical_approach_documentation

    if step == 4:
        form_id = session['form_data'].get('id')
        if form_id:
            existing_form = Forms.query.get(form_id)
            if existing_form and existing_form.user_id == current_user.id:
                form.past_performance.data = existing_form.past_performance

    return render_template('forms/submit_form.html', form=form, table_form=table_form, table_entries=table_entries, current_step=step, total_steps=total_steps)

@blueprint.route('/delete_form/<int:form_id>', methods=['POST'])
@login_required
def delete_form(form_id):
    form = Forms.query.get(form_id)
    if form and form.user_id == current_user.id:
        TableEntry.query.filter_by(form_id=form.id).delete()
        db.session.delete(form)
        db.session.commit()
        flash('Form and associated table entries deleted successfully.', 'success')
    else:
        flash('Form not found or you do not have permission to delete it.', 'error')
    return redirect(url_for('forms_blueprint.view_forms'))

# Table entry
@blueprint.route('/edit_entry/<int:entry_id>', methods=['GET', 'POST'])
@login_required
@proposal_charges_required
def edit_entry(entry_id):
    entry = TableEntry.query.get_or_404(entry_id)
    form = TableForm(obj=entry)

    if form.validate_on_submit():
        item_no = "{:03d}".format(form.item_no.data)
        entry.item_no = item_no
        entry.description = form.description.data
        entry.quantity = form.quantity.data
        entry.unit = form.unit.data
        entry.unit_price = form.unit_price.data
        db.session.commit()
        flash('Entry updated successfully.', 'success')
        return redirect(url_for('forms_blueprint.submit_form', step=2))

    return render_template('forms/edit_entry.html', form=form, entry=entry)

@blueprint.route('/delete_entry/<int:entry_id>', methods=['GET', 'POST'])
def delete_entry(entry_id):
    entry = TableEntry.query.get_or_404(entry_id)
    db.session.delete(entry)
    db.session.commit()
    flash('Entry deleted successfully.', 'success')
    return redirect(url_for('forms_blueprint.submit_form', step=2))

# Word document generation
@blueprint.route('/download_form/<int:form_id>', methods=['GET'])
@login_required
@proposal_charges_required
def download_form(form_id):
    form = Forms.query.get(form_id)
    
    if form and form.user_id == current_user.id:
        # Decrement the number of proposals if this is the first download
        if form.number_of_downloads == 0:
            if current_user.number_of_proposals > 0:
                current_user.number_of_proposals -= 1
            else:
                flash('You have no remaining proposals left to download.', 'error')
                return redirect(url_for('forms_blueprint.view_forms'))
        
        # Increment the download count, if it’s within the limit
        if form.number_of_downloads < 3:
            form.number_of_downloads += 1
            db.session.commit()
        else:
            flash('You have reached the maximum number of downloads for this form.', 'error')
            return redirect(url_for('forms_blueprint.view_forms'))
        
        # Generate the document
        document = Document()
        new_parser = HtmlToDocx()
        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run('Solicitation Number and Title')
        run.font.name = 'Arial'
        run.bold = True
        run.font.size = Pt(18)

        document.add_paragraph("\n")

        solicitation_number = document.add_paragraph()
        solicitation_number.alignment = WD_ALIGN_PARAGRAPH.CENTER
        solicitation_number_run = solicitation_number.add_run(str(form.solicitation_number))
        solicitation_number_run.bold = True
        solicitation_number_run.font.name = 'Arial'
        solicitation_number_run.font.size = Pt(16)

        document.add_paragraph("\n\n")

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(str(form.title))
        title_run.bold = True
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(16)

        def add_bold_paragraph(text):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(text)
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            return paragraph

        def add_normal_text(paragraph, text):
            run = paragraph.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(11)

        document.add_paragraph("\n\n\n\n\n\n\n\n\n")

        paragraph = add_bold_paragraph("Unique Entity ID: ")
        add_normal_text(paragraph, str(form.unique_entity_id))

        paragraph = add_bold_paragraph("Phone Number: ")
        add_normal_text(paragraph, str(form.phone_number))

        paragraph = add_bold_paragraph("POC Email: ")
        add_normal_text(paragraph, str(form.poc_email))

        paragraph = add_bold_paragraph("CAGE Code: ")
        add_normal_text(paragraph, str(form.cage_code))

        paragraph = add_bold_paragraph("EIN/GST-HST Number: ")
        add_normal_text(paragraph, str(form.ein_gst_hst_number))

        paragraph = add_bold_paragraph("POC: ")
        add_normal_text(paragraph, str(form.poc))

        document.add_page_break()

        table_entries = TableEntry.query.filter_by(form_id=form_id).all()

        if table_entries:
            table = document.add_table(rows=len(table_entries) + 1, cols=6)
            table.style = 'Table Grid'
            header = table.rows[0].cells
            header[0].text = 'Item No'
            header[1].text = 'Description'
            header[2].text = 'Quantity'
            header[3].text = 'Unit'
            header[4].text = 'Unit Price'
            header[5].text = 'Ext Price'

            for i, entry in enumerate(table_entries, start=1):
                row = table.rows[i].cells
                row[0].text = str(entry.item_no)
                row[1].text = entry.description
                row[2].text = str(entry.quantity)
                row[3].text = entry.unit
                row[4].text = f"${entry.unit_price:.2f}"
                row[5].text = f"${entry.quantity * entry.unit_price:.2f}"
        else:
            document.add_paragraph("No table entries found.")

        document.add_page_break()

        # Add Technical Approach Documentation
        
        technical_title = document.add_paragraph()
        technical_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        technical_title_run = technical_title.add_run('Technical Documentation')
        technical_title_run.bold = True
        technical_title_run.underline = True
        technical_title_run.font.name = 'Arial'
        technical_title_run.font.size = Pt(11)
        
        if form.technical_approach_documentation:
            new_parser.add_html_to_document(form.technical_approach_documentation, document)
        else:
            document.add_paragraph("No technical approach documentation provided.")

        document.add_page_break()

        # Add Past Performance
        past_title = document.add_paragraph()
        past_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        past_title_run = past_title.add_run('Past Performance')
        past_title_run.bold = True
        past_title_run.underline = True
        past_title_run.font.name = 'Arial'
        past_title_run.font.size = Pt(11)
        
        if form.past_performance:
            new_parser.add_html_to_document(form.past_performance, document)
        else:
            document.add_paragraph("No past performance information provided.")
            
        byte_io = BytesIO()
        document.save(byte_io)
        byte_io.seek(0)
        
        # Send the file
        response = send_file(
            byte_io, 
            as_attachment=True, 
            download_name='Completed_Proposal_Notice.docx', 
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Return the response and handle redirection with JavaScript in the frontend
        return response
    
    flash('Form not found or you do not have permission to download it.', 'error')
    return redirect(url_for('forms_blueprint.view_forms'))

# Errors

@login_manager.unauthorized_handler
def unauthorized_handler():
    return render_template('home/page-403.html'), 403


@blueprint.errorhandler(403)
def access_forbidden(error):
    return render_template('home/page-403.html'), 403


@blueprint.errorhandler(404)
def not_found_error(error):
    return render_template('home/page-404.html'), 404


@blueprint.errorhandler(500)
def internal_error(error):
    return render_template('home/page-500.html'), 500
